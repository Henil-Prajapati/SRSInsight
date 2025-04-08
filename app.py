from flask import Flask, render_template, request, jsonify, send_file
from tavily import TavilyClient
from groq import Groq
import json
import subprocess
import os
import tempfile
from langchain.memory import ConversationBufferMemory
import re
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
import base64

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Initialize memory
memory = ConversationBufferMemory()

# Get API keys from environment variables
TAVILY_API_KEY = os.getenv('TAVILY_API_KEY', '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')
GROQ_API_KEY = os.getenv('GROQ_API_KEY', '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')

# Initialize API clients
tavily_client = TavilyClient(api_key=TAVILY_API_KEY)
groq_client = Groq(api_key=GROQ_API_KEY)

# Helper functions (same as in your original code)
def extract_json(raw_output):
    match = re.search(r'\{.*\}', raw_output, re.DOTALL)
    if match:
        return match.group(0)
    raise ValueError("No valid JSON found in response")

def analyze_sentiment(app_description, groq_client):
    # ... (same as original)
    prompt = f"""
    Analyze the sentiment of this input: "{app_description}".
    Return:
    - Sentiment: Positive, Negative, Neutral, or Urgent
    - Score: 0.0 to 1.0 (confidence)
    Example: "Sentiment: Positive, Score: 0.85"
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=100
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error in sentiment analysis: {str(e)}"

def get_stakeholders(title, description, functionalities, groq_client):
    # ... (same as original)
    prompt = f"""
    You are a professional SRS developer in San Francisco with extensive experience crafting detailed specifications. Given:
    - App Title: "{title}"
    - Description: "{description}"
    - Desired Functionalities: "{functionalities}"

    Generate a stakeholder analysis for an SRS:
    1. Identify all relevant stakeholders with interest or influence in the project.
    2. For each, provide:
       - Name (e.g., "End-users")
       - Role: A detailed, multi-sentence description of their involvement, responsibilities, and interactions with the system, written in full narrative form.
       - Priority (High, Medium, Low): With a brief justification.
    3. Return *ONLY* a valid JSON object with no extra text outside the JSON:
       {{
         "stakeholders": [
           {{"name": "", "role": "", "priority": ""}}
         ]
       }}
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=1000
        )
        raw_output = response.choices[0].message.content.strip()
        cleaned_json = extract_json(raw_output)
        return json.loads(cleaned_json)
    except Exception as e:
        return {"stakeholders": []}

def get_srs(title, description, functionalities, stakeholders_json, web_output, sentiment, groq_client):
    # ... (same as original)
    stakeholders = json.dumps(stakeholders_json)
    prompt = f"""
    You are a professional SRS developer in San Francisco with extensive experience crafting detailed specifications. Given:
    - App Title: "{title}"
    - Description: "{description}"
    - Desired Functionalities: "{functionalities}"
    - Stakeholders: {stakeholders}
    - Web Search Results: "{web_output}"
    - Sentiment: "{sentiment}"

    Generate a complete SRS per IEEE 29148, excluding diagrams. Provide detailed, multi-sentence, narrative-style descriptions for each section in full paragraphs, except for Specific Requirements, which should be numbered items. Include:
    1. **Introduction**:
       - Purpose: A detailed paragraph explaining the app's intent, benefits, and target audience.
       - Scope: A paragraph detailing what the app will achieve, its boundaries, and exclusions, with examples.
       - Definitions: A paragraph defining key terms with thorough explanations.
       - References: A placeholder paragraph as "TBD" with a note on future updates.
    2. **Overall Description**:
       - Product Perspective: A paragraph on context within its ecosystem, interfaces, and novelty.
       - User Needs: A paragraph detailing primary user needs tied to functionalities, with examples.
       - Assumptions/Dependencies: A paragraph on external factors, with justifications.
    3. **Specific Requirements**:
       - Functional: A list of numbered items, each a detailed paragraph (3-5 sentences) describing one functionality, including inputs, outputs, and conditions.
       - Non-Functional: A list of numbered items, each a detailed paragraph (2-3 sentences) covering performance, security, usability, etc., with metrics or examples.
       - Design Constraints: A list of numbered items, each a detailed paragraph (2-3 sentences) detailing technical or operational limits with rationale.
    4. **Stakeholder Analysis**: Use provided stakeholders' detailed roles as-is.
    5. **Risk Analysis**: A list of risks, each with a detailed risk paragraph and a mitigation paragraph.

    Return *ONLY* a valid JSON object with no extra text outside the JSON:
    {{
      "introduction": {{"purpose": "", "scope": "", "definitions": "", "references": ""}},
      "overall_description": {{"product_perspective": "", "user_needs": "", "assumptions_dependencies": ""}},
      "specific_requirements": {{"functional": [""], "non_functional": [""], "design_constraints": [""]}},
      "stakeholder_analysis": [],
      "risk_analysis": [{{"risk": "", "mitigation": ""}}]
    }}
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=4000
        )
        raw_output = response.choices[0].message.content.strip()
        cleaned_json = extract_json(raw_output)
        return json.loads(cleaned_json)
    except Exception as e:
        return {}

def generate_docx(srs_json):
    # ... (same as original)
    doc = Document()
    doc.add_heading(f"{srs_json['introduction']['purpose'].split('.')[0]} SRS", 0)
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("1.1 Purpose")
    doc.add_paragraph(srs_json['introduction']['purpose'])
    doc.add_paragraph("1.2 Scope")
    doc.add_paragraph(srs_json['introduction']['scope'])
    doc.add_paragraph("1.3 Definitions")
    doc.add_paragraph(srs_json['introduction']['definitions'])
    doc.add_paragraph("1.4 References")
    doc.add_paragraph(srs_json['introduction']['references'])
    doc.add_heading("2. Overall Description", level=1)
    doc.add_paragraph("2.1 Product Perspective")
    doc.add_paragraph(srs_json['overall_description']['product_perspective'])
    doc.add_paragraph("2.2 User Needs")
    doc.add_paragraph(srs_json['overall_description']['user_needs'])
    doc.add_paragraph("2.3 Assumptions and Dependencies")
    doc.add_paragraph(srs_json['overall_description']['assumptions_dependencies'])
    doc.add_heading("3. Specific Requirements", level=1)
    doc.add_paragraph("3.1 Functional Requirements")
    for i, req in enumerate(srs_json['specific_requirements']['functional'], 1):
        doc.add_paragraph(f"3.1.{i} {req.split('.')[0]}")
        doc.add_paragraph(req)
    doc.add_paragraph("3.2 Non-Functional Requirements")
    for i, req in enumerate(srs_json['specific_requirements']['non_functional'], 1):
        doc.add_paragraph(f"3.2.{i} {req.split('.')[0]}")
        doc.add_paragraph(req)
    doc.add_paragraph("3.3 Design Constraints")
    for i, req in enumerate(srs_json['specific_requirements']['design_constraints'], 1):
        doc.add_paragraph(f"3.3.{i} {req.split('.')[0]}")
        doc.add_paragraph(req)
    doc.add_heading("4. Stakeholder Analysis", level=1)
    for i, stakeholder in enumerate(srs_json['stakeholder_analysis'], 1):
        doc.add_paragraph(f"4.{i} {stakeholder['name']} ({stakeholder['priority']})")
        doc.add_paragraph(stakeholder['role'])
    doc.add_heading("5. Risk Analysis", level=1)
    for i, risk in enumerate(srs_json['risk_analysis'], 1):
        doc.add_paragraph(f"5.{i} Risk")
        doc.add_paragraph(risk['risk'])
        doc.add_paragraph(f"5.{i}.1 Mitigation")
        doc.add_paragraph(risk['mitigation'])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_use_case_plantuml_code(app_description, stakeholders, groq_client):
    # ... (same as original)
    prompt = f"""
    Given the app description and stakeholders, generate the PlantUML code for a use case diagram. 
    App description: "{app_description}"
    Stakeholders: {stakeholders}
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=500
        )
        content = response.choices[0].message.content
        if "```" in content:
            content = content.replace("```plantuml", "").replace("```", "").strip()
        if "@startuml" not in content:
            content = "@startuml\n" + content
        if "@enduml" not in content:
            content += "\n@enduml"
        return content
    except Exception as e:
        return None

def generate_activity_plantuml_code(app_description, groq_client):
    # ... (same as original)
    prompt = f"""
    You are a software architect. Generate the PlantUML code for an activity diagram that models the primary flow of the application based on the description below.

    App Description: "{app_description}"

    The diagram should follow proper PlantUML syntax using @startuml and @enduml. Only return the PlantUML code without any explanations or additional text.
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=500
        )
        content = response.choices[0].message.content
        if "```" in content:
            content = content.replace("```plantuml", "").replace("```", "").strip()
        if "@startuml" not in content:
            content = "@startuml\n" + content
        if "@enduml" not in content:
            content += "\n@enduml"
        return content
    except Exception as e:
        return None

def generate_class_plantuml_code(app_description, stakeholders, groq_client):
    # ... (same as original)
    prompt = f"""
    Generate a PlantUML class diagram for this application.
    App description: "{app_description}"
    Stakeholders: {stakeholders}
    """
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=1500
        )
        content = response.choices[0].message.content
        if "```" in content:
            content = content.replace("```plantuml", "").replace("```", "").strip()
        if "@startuml" not in content:
            content = "@startuml\nskinparam classAttributeIconSize 0\n" + content
        if "@enduml" not in content:
            content += "\n@enduml"
        return content
    except Exception as e:
        return None

def generate_uml_image(plantuml_code):
    try:
        if not plantuml_code:
            print("No PlantUML code provided")
            return None
            
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, "diagram.puml")
            with open(temp_file_path, "w") as f:
                f.write(plantuml_code)
                
            # Use the plantuml.jar in the current directory
            plantuml_jar_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "plantuml.jar")
            
            if not os.path.exists(plantuml_jar_path):
                print(f"PlantUML jar not found at: {plantuml_jar_path}")
                return None
                
            print(f"Generating diagram using PlantUML jar at: {plantuml_jar_path}")
            result = subprocess.run(
                ["java", "-jar", plantuml_jar_path, "-verbose", temp_file_path],
                capture_output=True,
                text=True,
                cwd=temp_dir
            )
            
            if result.returncode != 0:
                print(f"PlantUML generation failed with error: {result.stderr}")
                return None
                
            expected_output = os.path.join(temp_dir, "diagram.png")
            if not os.path.exists(expected_output):
                print(f"Expected output file not found at: {expected_output}")
                return None
                
            with open(expected_output, "rb") as f:
                image_data = f.read()
                print("Successfully generated diagram")
                return image_data
                
    except Exception as e:
        print(f"Error generating UML image: {str(e)}")
        return None

def process_app_description(title, description, functionalities, diagram_type, tavily_client, groq_client):
    memory.save_context({"input": f"{title}: {description}, {functionalities}"}, {"output": "Processing..."})
    search_query = f"functionalities of existing systems similar to '{description}'"
    search_results = tavily_client.search(query=search_query, max_results=5, search_depth="basic")
    web_output = "\n".join([result.get("content", "") for result in search_results.get("results", [])])
    sentiment = analyze_sentiment(description, groq_client)
    stakeholders_json = get_stakeholders(title, description, functionalities, groq_client)
    stakeholders_str = "\n".join([f"{s['name']} ({s['priority']}): {s['role']}" for s in stakeholders_json['stakeholders']])
    srs_json = get_srs(title, description, functionalities, stakeholders_json, web_output, sentiment, groq_client)
    docx_file = generate_docx(srs_json)
    
    # Generate diagram
    diagram_html = ""
    if diagram_type == "Use Case Diagram":
        plantuml_code = generate_use_case_plantuml_code(f"{title}: {description}", stakeholders_str, groq_client)
        if plantuml_code:
            image_bytes = generate_uml_image(plantuml_code)
            if image_bytes:
                # Convert bytes to base64 string
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                diagram_html = f'<h3>Use Case Diagram</h3><pre>{plantuml_code}</pre><img src="data:image/png;base64,{image_base64}">'
    elif diagram_type == "Activity Diagram":
        plantuml_code = generate_activity_plantuml_code(f"{title}: {description}", groq_client)
        if plantuml_code:
            image_bytes = generate_uml_image(plantuml_code)
            if image_bytes:
                # Convert bytes to base64 string
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                diagram_html = f'<h3>Activity Diagram</h3><pre>{plantuml_code}</pre><img src="data:image/png;base64,{image_base64}">'
    elif diagram_type == "Class Diagram":
        plantuml_code = generate_class_plantuml_code(f"{title}: {description}", stakeholders_str, groq_client)
        if plantuml_code:
            image_bytes = generate_uml_image(plantuml_code)
            if image_bytes:
                # Convert bytes to base64 string
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                diagram_html = f'<h3>Class Diagram</h3><pre>{plantuml_code}</pre><img src="data:image/png;base64,{image_base64}">'

    memory.save_context({"input": f"{title}: {description}, {functionalities}"}, {"output": json.dumps(srs_json)})
    return srs_json, stakeholders_str, docx_file, diagram_html

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        required_fields = ['title', 'description', 'functionalities', 'diagramType']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400

        title = data['title']
        description = data['description']
        functionalities = data['functionalities']
        diagram_type = data['diagramType']

        try:
            srs_json, stakeholders_str, docx_file, diagram_html = process_app_description(
                title, description, functionalities, diagram_type, tavily_client, groq_client
            )
            
            # Save the document
            doc_filename = f"{title}_SRS.docx"
            with open(doc_filename, "wb") as f:
                f.write(docx_file.getvalue())
            
            return jsonify({
                'title': title,
                'srs_json': srs_json,
                'stakeholders_str': stakeholders_str,
                'diagram_html': diagram_html
            })
        except Exception as e:
            return jsonify({'error': f'Processing error: {str(e)}'}), 500

    except Exception as e:
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/download/<filename>')
def download(filename):
    try:
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({'error': f'Download error: {str(e)}'}), 404

if __name__ == '__main__':
    app.run(debug=True)